import os
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def read_transcript(file_path):
    """Reads raw unstructured transcript data."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Error: The file '{file_path}' was not found.")
    with open(file_path, "r", encoding="utf-8") as file:
        return file.read()

def query_local_llm(transcript_content):
    """Connects to local LM Studio server and executes structured prompting."""
    # Initialize the client pointing to LM Studio's local server port
    client = OpenAI(base_url="http://localhost:1234/v1", api_key="lm-studio")
    
    # Constructing Microsoft's 4-part Framework Prompt
    # 1. Goal: "Summarize the attached transcript..."
    # 2. Context: "...for an official incident report."
    # 3. Source: "Use only the provided text."
    # 4. Expectation: "Return three sections: Issue, Troubleshooting Steps, and Resolution."
    system_prompt = (
        "You are an expert technical operations writer. Your task is to summarize the attached transcript "
        "for an official corporate incident report. Use ONLY the provided text source to extract facts—do not "
        "extrapolate or invent details. Your output MUST return exactly three distinct sections labeled precisely as: "
        "### Issue\n"
        "### Troubleshooting Steps\n"
        "### Resolution\n"
        "Keep the language professional, concise, and technical."
    )
    
    print("Sending transcript to local LLM inside LM Studio...")
    
    try:
        response = client.chat.completions.create(
            model="local-model", # LM Studio automatically overrides this with whatever model you have loaded
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": f"Here is the transcript data:\n\n{transcript_content}"}
            ],
            temperature=0.2 # Lower temperature guarantees factual accuracy based strictly on the source
        )
        return response.choices[0].message.content
    except Exception as e:
        raise ConnectionError(f"Could not connect to LM Studio. Ensure the server is running on port 1234. Details: {e}")

def create_word_document(llm_text, output_path):
    """Compiles text strings into a professional Microsoft Word document layout."""
    print("Generating formatted Word Document...")
    doc = Document()
    
    # Set standard margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Document Header/Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run("INCIDENT REPORT: PRODUCTION DATABASE CRASH")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    
    # Add a thin separator line representation
    doc.add_paragraph("─" * 55)

    # Parsing the LLM markdown response to apply proper Word styling structures
    lines = llm_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Check for section boundaries delivered by LLM
        if line.startswith("### Issue") or line.lower().startswith("issue:"):
            p = doc.add_paragraph()
            run = p.add_run("1. Incident Summary / Issue")
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            
elif line.startswith("### Troubleshooting Steps") or line.lower().startswith("troubleshooting steps:"):
            p = doc.add_paragraph()
            run = p.add_run("2. Diagnostic & Troubleshooting Actions")
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            
        elif line.startswith("### Resolution") or line.lower().startswith("resolution:"):
            p = doc.add_paragraph()
            run = p.add_run("3. Final Resolution & Post-Mortem Tasks")
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            
        else:
            # Clean up inline markdown bold markers if the LLM generated them
            clean_line = line.replace("**", "").replace("* ", "• ")
            p = doc.add_paragraph()
            run = p.add_run(clean_line)
            run.font.name = 'Arial'
            run.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.15

    # Ensure output directory exists and save document
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)
    print(f"Success! Document generated and saved to: {output_path}")

if __name__ == "__main__":
    TRANSCRIPT_FILE = "data/chat_transcript.txt"
    OUTPUT_DOCX = "outputs/Incident_Report.docx"
    
    try:
        raw_data = read_transcript(TRANSCRIPT_FILE)
        summary_text = query_local_llm(raw_data)
        create_word_document(summary_text, OUTPUT_DOCX)
    except Exception as error:
        print(f"\nPipeline Execution Failed: {error}")